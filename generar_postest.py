from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ============ HEADER ============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CUESTIONARIO POST-TEST')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0xBE, 0x12, 0x3C)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Evaluaci\u00f3n de Experiencia de Usuario \u2014 AnaVictoria Flores')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x78, 0x71, 0x6C)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Responde en funci\u00f3n de tu experiencia DESPU\u00c9S de haber navegado y probado el sitio web.')
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0xA8, 0xA2, 0x9E)

doc.add_paragraph()

# ============ HELPERS ============
def add_section_title(doc, number, title):
    p = doc.add_paragraph()
    run = p.add_run(f'{number}. {title}')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '6',
        qn('w:space'): '4', qn('w:color'): 'E2D5CE'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_question(doc, number, text, hint=''):
    p = doc.add_paragraph()
    run = p.add_run(f'{number}. {text}')
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    if hint:
        run2 = p.add_run(f'  ({hint})')
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0xA8, 0xA2, 0x9E)
        run2.italic = True

def add_option(doc, letter, text):
    p = doc.add_paragraph()
    run = p.add_run(f'____ {letter}. {text}')
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)

def add_open_question(doc, number, text):
    p = doc.add_paragraph()
    run = p.add_run(f'{number}. {text}')
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(10)
    for _ in range(3):
        doc.add_paragraph('')

def add_single_select(doc, number, text, options):
    add_question(doc, number, text)
    for i, opt in enumerate(options):
        letter = chr(97 + i)
        add_option(doc, letter, opt)

def add_multi_select(doc, number, text, options):
    add_question(doc, number, text, 'Puede seleccionar m\u00e1s de una')
    for i, opt in enumerate(options):
        letter = chr(97 + i)
        add_option(doc, letter, opt)

# ================================================================
# SECTION A
# ================================================================
add_section_title(doc, 'A', 'SATISFACCI\u00d3N GENERAL')

add_single_select(doc, 'A1', 'En general, \u00bfqu\u00e9 tan satisfecho(a) quedaste con tu experiencia en el sitio web de AnaVictoria?', [
    'Muy insatisfecho', 'Insatisfecho', 'Neutral', 'Satisfecho', 'Muy satisfecho'
])
add_single_select(doc, 'A2', '\u00bfEl dise\u00f1o visual y la est\u00e9tica del sitio reflejan la calidad que esperar\u00edas de una florister\u00eda?', [
    'No, para nada', 'Poco', 'Neutral', 'Bastante', 'Totalmente'
])
add_single_select(doc, 'A3', '\u00bfQu\u00e9 tan probable es que recomiendes AnaVictoria a un amigo o familiar?', [
    'Nada probable', 'Poco probable', 'Neutral', 'Probable', 'Muy probable'
])

# ================================================================
# SECTION B
# ================================================================
add_section_title(doc, 'B', 'USABILIDAD (ESCALA SUS)')

p = doc.add_paragraph()
run = p.add_run('Indica tu nivel de acuerdo con cada afirmaci\u00f3n (1 = Totalmente en desacuerdo, 5 = Totalmente de acuerdo):')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x78, 0x71, 0x6C)

sus_questions = [
    'B1. Creo que me gustar\u00eda usar este sitio web con frecuencia',
    'B2. Encontr\u00e9 el sitio innecesariamente complejo',
    'B3. Pens\u00e9 que el sitio era f\u00e1cil de usar',
    'B4. Creo que necesitar\u00eda ayuda t\u00e9cnica para usar este sitio',
    'B5. Las funciones del sitio est\u00e1n bien integradas',
    'B6. Hay demasiada inconsistencia en el sitio',
    'B7. Imagino que la mayor\u00eda de personas aprender\u00eda a usar este sitio muy r\u00e1pido',
    'B8. Encontr\u00e9 el sitio muy dif\u00edcil de navegar',
    'B9. Me sent\u00ed muy confiado(a) navegando por el sitio',
    'B10. Necesit\u00e9 aprender muchas cosas antes de poder usar bien el sitio',
]

for q in sus_questions:
    add_question(doc, '', q)
    doc.add_paragraph('    1 _____   2 _____   3 _____   4 _____   5 _____')

# ================================================================
# SECTION C
# ================================================================
add_section_title(doc, 'C', 'NAVEGACI\u00d3N Y ESTRUCTURA')

add_single_select(doc, 'C1', 'Encontrar lo que buscabas dentro del sitio fue:', [
    'Muy dif\u00edcil', 'Dif\u00edcil', 'Aceptable', 'F\u00e1cil', 'Muy f\u00e1cil'
])
add_single_select(doc, 'C2', '\u00bfLa barra de navegaci\u00f3n superior (Inicio, Ramos, Detalles Florales, Pedidos Especiales) te result\u00f3 clara y \u00fatil?', [
    'Nada clara', 'Poco clara', 'Neutral', 'Clara', 'Muy clara'
])
add_multi_select(doc, 'C3', '\u00bfQu\u00e9 secciones del sitio visitaste o exploraste?', [
    'P\u00e1gina de inicio', 'Ramos de Flores', 'Detalles Florales',
    'P\u00e1gina de detalle de producto', 'Carrito de compras',
    'Pedidos Especiales (formulario)', 'Bot\u00f3n / chat de WhatsApp'
])
add_single_select(doc, 'C4', '\u00bfLa p\u00e1gina de inicio te ayud\u00f3 a entender r\u00e1pidamente qu\u00e9 ofrece AnaVictoria?', [
    'No, fue confusa', 'Poco', 'Neutral', 'S\u00ed, bastante', 'S\u00ed, completamente'
])
add_single_select(doc, 'C5', '\u00bfEl selector de ocasiones (Cumplea\u00f1os, Aniversario, Bodas, etc.) en la p\u00e1gina de inicio te pareci\u00f3 \u00fatil?', [
    'S\u00ed, muy \u00fatil', 'Poco \u00fatil', 'No lo not\u00e9', 'No me sirvi\u00f3'
])

# ================================================================
# SECTION D
# ================================================================
add_section_title(doc, 'D', 'CAT\u00c1LOGO Y PRODUCTOS')

add_single_select(doc, 'D1', '\u00bfC\u00f3mo calificas la calidad y claridad de las fotograf\u00edas de los productos?', [
    'Muy malas', 'Malas', 'Aceptables', 'Buenas', 'Excelentes'
])
add_single_select(doc, 'D2', '\u00bfLa informaci\u00f3n mostrada en cada producto fue suficiente para tomar una decisi\u00f3n de compra?', [
    'Insuficiente', 'Poco suficiente', 'Neutral', 'Suficiente', 'M\u00e1s que suficiente'
])
add_single_select(doc, 'D3', '\u00bfUsaste los filtros por ocasi\u00f3n o el ordenamiento (por precio, valoraci\u00f3n)?', [
    'S\u00ed, los us\u00e9 y funcionaron bien', 'S\u00ed, los us\u00e9 pero no funcionaron como esperaba',
    'No los us\u00e9', 'No not\u00e9 que hubiera filtros'
])
add_multi_select(doc, 'D4', 'En la p\u00e1gina de detalle del producto, \u00bfqu\u00e9 elementos te parecieron m\u00e1s \u00fatiles? (Selecciona hasta 3)', [
    'La imagen grande del producto', 'Las estrellas de valoraci\u00f3n y rese\u00f1as',
    'Las etiquetas (ej. "Rom\u00e1ntico", "Blanco")', 'El precio con nota de empaque artesanal',
    'La descripci\u00f3n del producto', 'Las garant\u00edas (env\u00edo gratis, frescura garantizada)',
    'Productos relacionados ("Tambi\u00e9n te puede gustar")'
])
add_single_select(doc, 'D5', '\u00bfSentiste que te falt\u00f3 informaci\u00f3n visual para calcular el tama\u00f1o real de los arreglos?', [
    'S\u00ed, me falt\u00f3 mucho', 'Me falt\u00f3 un poco', 'Neutral', 'Casi suficiente', 'No me falt\u00f3 nada'
])

# ================================================================
# SECTION E
# ================================================================
add_section_title(doc, 'E', 'CARRITO Y PROCESO DE COMPRA')

add_single_select(doc, 'E1', '\u00bfAgregaste alg\u00fan producto al carrito de compras?', ['S\u00ed', 'No'])
add_single_select(doc, 'E2', 'Si agregaste productos, \u00bfc\u00f3mo calificas la experiencia del carrito?', [
    'Muy mala', 'Mala', 'Aceptable', 'Buena', 'Excelente'
])
add_single_select(doc, 'E3', '\u00bfEl resumen del pedido (subtotal, IVA 15%, env\u00edo, total) te pareci\u00f3 claro?', [
    'Nada claro', 'Poco claro', 'Neutral', 'Claro', 'Muy claro'
])
add_single_select(doc, 'E4', '\u00bfLlegaste hasta la pantalla de confirmaci\u00f3n con los datos bancarios?', [
    'S\u00ed, vi los datos de pago', 'No llegu\u00e9 hasta ese paso'
])
add_single_select(doc, 'E5', '\u00bfQu\u00e9 te pareci\u00f3 el flujo de pago (transferencia + comprobante por WhatsApp)?', [
    'Me pareci\u00f3 confiable y claro', 'Me gener\u00f3 inseguridad',
    'Complicado tener que ir a WhatsApp aparte', 'R\u00e1pido y sencillo',
    'No tengo opini\u00f3n / no llegu\u00e9 a ese paso'
])
add_single_select(doc, 'E6', '\u00bfTe molest\u00f3 que no se pidieran datos de env\u00edo en el carrito?', [
    'S\u00ed, deber\u00eda pedirlos ah\u00ed', 'No, est\u00e1 bien que se coordine por WhatsApp', 'Me es indiferente'
])

# ================================================================
# SECTION F
# ================================================================
add_section_title(doc, 'F', 'PEDIDOS ESPECIALES (FORMULARIO)')

add_single_select(doc, 'F1', '\u00bfViste o interactuaste con el formulario de Pedidos Especiales?', ['S\u00ed', 'No'])
add_single_select(doc, 'F2', 'Si lo viste, \u00bfel formulario te pareci\u00f3:', [
    'Muy largo / tedioso', 'Algo largo', 'Adecuado', 'Bien organizado', 'Muy completo y f\u00e1cil de llenar'
])
add_single_select(doc, 'F3', '\u00bfQu\u00e9 tan claro te pareci\u00f3 el paso a paso al lado del formulario?', [
    'Confuso', 'Poco claro', 'Neutral', 'Claro', 'Muy claro'
])

# ================================================================
# SECTION G
# ================================================================
add_section_title(doc, 'G', 'WHATSAPP Y COMUNICACI\u00d3N')

add_single_select(doc, 'G1', '\u00bfNotaste el bot\u00f3n flotante de WhatsApp?', ['S\u00ed', 'No'])
add_single_select(doc, 'G2', '\u00bfUsar\u00edas WhatsApp como canal principal para dudas o entregas?', [
    'S\u00ed, me parece pr\u00e1ctico', 'No, preferir\u00eda otro medio', 'Me da igual'
])
add_single_select(doc, 'G3', '\u00bfTe gustar\u00eda recibir notificaciones del pedido por WhatsApp?', [
    'S\u00ed', 'No', 'Prefiero por correo electr\u00f3nico'
])

# ================================================================
# SECTION H
# ================================================================
add_section_title(doc, 'H', 'EXPERIENCIA M\u00d3VIL')

add_single_select(doc, 'H1', '\u00bfProbaste el sitio en un tel\u00e9fono celular?', ['S\u00ed', 'No, solo en computadora'])
add_single_select(doc, 'H2', 'Si lo probaste en celular, \u00bfc\u00f3mo calificas la experiencia m\u00f3vil?', [
    'Muy mala', 'Mala', 'Aceptable', 'Buena', 'Excelente'
])
add_multi_select(doc, 'H3', 'En el pre-test mencionaste problemas de dise\u00f1o en m\u00f3vil. \u00bfSe presentaron en AnaVictoria?', [
    'Textos peque\u00f1os o dif\u00edciles de leer', 'Im\u00e1genes que no se pueden ampliar',
    'Botones peque\u00f1os o dif\u00edciles de presionar', 'Ventanas emergentes que tapan la pantalla',
    'No tuve ninguno de esos problemas'
])

# ================================================================
# SECTION I
# ================================================================
add_section_title(doc, 'I', 'CONFIANZA Y SEGURIDAD')

add_single_select(doc, 'I1', '\u00bfEl sitio te transmiti\u00f3 confianza para ingresar tus datos?', [
    'Nada de confianza', 'Poca confianza', 'Neutral', 'Confianza', 'Mucha confianza'
])
add_multi_select(doc, 'I2', '\u00bfQu\u00e9 informaci\u00f3n de respaldo te dio m\u00e1s confianza? (Hasta 2)', [
    'Los testimonios de clientes', 'La garant\u00eda de frescura',
    'Tener WhatsApp visible siempre', 'Datos de contacto en el footer',
    'Zonas de entrega detalladas', 'Fotos reales de los arreglos', 'Ninguno'
])
add_open_question(doc, 'I3', '\u00bfConsideras que falt\u00f3 alguna pol\u00edtica o sello de confianza en el sitio?')

# ================================================================
# SECTION J
# ================================================================
add_section_title(doc, 'J', 'PREGUNTAS ABIERTAS')

add_open_question(doc, 'J1', 'Describe tu MEJOR experiencia al usar el sitio. \u00bfQu\u00e9 te gust\u00f3 m\u00e1s?')
add_open_question(doc, 'J2', 'Describe tu PEOR experiencia o frustraci\u00f3n. \u00bfQu\u00e9 cambiar\u00edas?')
add_open_question(doc, 'J3', 'Si pudieras cambiar UNA SOLA COSA del sitio, \u00bfqu\u00e9 ser\u00eda?')
add_open_question(doc, 'J4', '\u00bfHay alguna funcionalidad que esperabas y no viste?')
add_open_question(doc, 'J5', '\u00bfRecomendar\u00edas este sitio a alguien m\u00e1s? \u00bfPor qu\u00e9?')

p = doc.add_paragraph()
run = p.add_run('\u2500' * 55)
run.font.color.rgb = RGBColor(0xE2, 0xD5, 0xCE)

add_open_question(doc, 'J6', '\u00bfC\u00f3mo se compara tu experiencia en AnaVictoria con tu proceso de compra habitual? \u00bfFue similar, m\u00e1s f\u00e1cil, m\u00e1s dif\u00edcil?')

# ================================================================
# SECTION K
# ================================================================
add_section_title(doc, 'K', 'COMPARATIVA PRE-TEST VS. POST-TEST')

add_single_select(doc, 'K1', '\u00bfEl sitio te permiti\u00f3 buscar/filtrar por tus criterios importantes del pre-test?', [
    'S\u00ed, completamente', 'Solo parcialmente', 'No, no encontr\u00e9 esa opci\u00f3n'
])

add_question(doc, 'K2', 'Califica a AnaVictoria en cada aspecto seg\u00fan tus prioridades del pre-test:')

aspects = ['Velocidad de carga', 'Claridad y calidad de fotograf\u00edas', 'Pocos pasos para completar el pago', 'Chat de soporte visible y accesible']
for asp in aspects:
    p = doc.add_paragraph()
    run = p.add_run(f'  {asp}:   1 ___   2 ___   3 ___   4 ___   5 ___')
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Cm(0.5)

add_open_question(doc, 'K3', '\u00bfEl bot\u00f3n "Confirmar y ver datos de pago" te pareci\u00f3 adecuado vs. tu frase preferida del pre-test?')
add_open_question(doc, 'K4', '\u00bfC\u00f3mo fue descubrir AnaVictoria vs. lo que esperabas del pre-test?')

# ================================================================
# SECTION L
# ================================================================
add_section_title(doc, 'L', 'PARA FINALIZAR')

add_open_question(doc, 'L1', 'En una palabra, \u00bfc\u00f3mo describir\u00edas tu experiencia general?')
add_open_question(doc, 'L2', '\u00bfHay algo m\u00e1s que quieras agregar?')
add_single_select(doc, 'L3', '\u00bfTe gustar\u00eda recibir una copia de los resultados?', ['S\u00ed', 'No'])

p = doc.add_paragraph()
run = p.add_run('Si respondiste que s\u00ed, ingresa tu correo electr\u00f3nico (opcional):')
run.font.size = Pt(10.5)
doc.add_paragraph('Correo: _______________________________________________')

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Muchas gracias por tu participaci\u00f3n. Tus respuestas nos ayudar\u00e1n a mejorar la experiencia de AnaVictoria.')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0xA8, 0xA2, 0x9E)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Este cuestionario es an\u00f3nimo y con fines acad\u00e9micos. Ning\u00fan dato personal ser\u00e1 compartido con terceros.')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xA8, 0xA2, 0x9E)
run.italic = True

output_path = r'C:\Users\Redbl\source\repos\Version con IA\Post-Test - Cuestionario AnaVictoria.docx'
doc.save(output_path)
print(f'OK: {output_path}')
